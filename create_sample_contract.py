"""
Generates sample_contract.docx: a small but *structurally real* stand-in for
the 60-page contract in the problem statement.

It deliberately contains the exact failure mode described in the brief:
  - A price is DEFINED once in Section 4.2 (wrapped in a Word bookmark).
  - The SAME price is quoted again in the Recitals ("page 1") and in the
    Signature Summary ("page 56") via REF fields pointing at that bookmark,
    not via copy-pasted text.
  - A pricing table computes Line Total = Qty * Unit Price, and a Grand
    Total = SUM(line totals), using real Word formula fields.

This is what "page 22 and page 56 don't update" looks like in the actual
file format: two REF fields with stale cached results, and a SUM field
whose cached result no longer matches its operands.
"""
import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_bookmark(paragraph, name, bm_id, text, bold=False):
    """Insert `text` as a run wrapped in a named bookmark inside `paragraph`."""
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(bm_id))
    start.set(qn('w:name'), name)
    paragraph._p.append(start)

    run = paragraph.add_run(text)
    run.bold = bold

    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bm_id))
    paragraph._p.append(end)
    return run


def add_ref_field(paragraph, bookmark_name, cached_text):
    """
    Insert a real Word REF field: { REF bookmark_name \\* MERGEFORMAT }
    with `cached_text` as the on-disk cached result (what shows before
    the user presses F9 / Update Fields). This cached text is exactly
    the thing that goes stale in every editor that doesn't track it.
    """
    r1 = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    r1._r.append(fld_begin)

    r2 = paragraph.add_run()
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = f' REF {bookmark_name} \\* MERGEFORMAT '
    r2._r.append(instr)

    r3 = paragraph.add_run()
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    r3._r.append(fld_sep)

    r4 = paragraph.add_run(cached_text)  # cached/displayed result

    r5 = paragraph.add_run()
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r5._r.append(fld_end)


def add_formula_cell(cell, formula, cached_text):
    """Put a real Word formula field (e.g. =B2*C2 or =SUM(ABOVE)) in a table cell."""
    p = cell.paragraphs[0]
    r1 = p.add_run()
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
    r1._r.append(b)
    r2 = p.add_run()
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = f' ={formula} \\# "$#,##0.00" '
    r2._r.append(instr)
    r3 = p.add_run()
    sep = OxmlElement('w:fldChar'); sep.set(qn('w:fldCharType'), 'separate')
    r3._r.append(sep)
    p.add_run(cached_text)
    r5 = p.add_run()
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
    r5._r.append(e)


def build():
    d = docx.Document()

    d.add_heading('Master Services Agreement', level=0)

    d.add_heading('Recitals', level=1)
    p = d.add_paragraph('This Agreement is entered into with a total Purchase Price of ')
    add_ref_field(p, 'clause_4_2_price', '$120,000.00')
    p.add_run('.  ("page 1")')

    d.add_heading('Section 4 — Payment Terms', level=1)
    d.add_paragraph('4.1 Payment shall be made in accordance with the schedule below.')
    p42 = d.add_paragraph('4.2 Purchase Price. The total Purchase Price payable by Client is ')
    add_bookmark(p42, 'clause_4_2_price', 100, '$120,000.00', bold=True)
    p42.add_run('.  ("page 22")')

    d.add_heading('Schedule A — Pricing', level=1)
    table = d.add_table(rows=4, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = 'Item', 'Qty', 'Line Total'

    rows_data = [
        ('Implementation', '1', 'B2*80000', '$80,000.00'),
        ('Support (annual)', '2', 'B3*20000', '$40,000.00'),
    ]
    for i, (item, qty, formula, cached) in enumerate(rows_data, start=1):
        cells = table.rows[i].cells
        cells[0].text = item
        cells[1].text = qty
        add_formula_cell(cells[2], formula, cached)

    total_cells = table.rows[3].cells
    total_cells[0].text = 'Grand Total'
    total_cells[1].text = ''
    add_formula_cell(total_cells[2], 'SUM(ABOVE)', '$120,000.00')

    d.add_heading('Signature Summary', level=1)
    p56 = d.add_paragraph('By signing below, the parties agree to a Purchase Price of ')
    add_ref_field(p56, 'clause_4_2_price', '$120,000.00')
    p56.add_run(', payable per Schedule A.  ("page 56")')

    d.save('sample_contract.docx')
    print('wrote sample_contract.docx')


if __name__ == '__main__':
    build()
